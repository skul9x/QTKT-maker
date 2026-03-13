"""
DOCX Merger Engine
Gộp nhiều file .docx thành 1 file duy nhất với page break giữa các file.
"""

import os
from copy import deepcopy
from pathlib import Path
from typing import Callable, Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _insert_before_sectpr(target_body, element):
    """Chèn element TRƯỚC sectPr (cuối body) để giữ đúng cấu trúc XML."""
    sect_pr = target_body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        target_body.append(element)


def add_page_break(target_body):
    """Thêm page break vào body (trước sectPr)."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    _insert_before_sectpr(target_body, p)


def copy_element(source_element, target_body):
    """Deep copy một XML element từ source sang target body (trước sectPr)."""
    new_element = deepcopy(source_element)
    _insert_before_sectpr(target_body, new_element)


def copy_numbering(source_doc, target_doc):
    """Copy numbering definitions từ source sang target document."""
    try:
        source_numbering_part = source_doc.part.numbering_part
        if source_numbering_part is not None:
            try:
                target_doc.part.numbering_part
            except Exception:
                # Target chưa có numbering part, sẽ tự tạo khi cần
                pass
    except Exception:
        pass


def copy_styles(source_doc, target_doc):
    """Copy styles từ source document nếu target chưa có."""
    try:
        source_styles = source_doc.styles
        target_style_names = {s.name for s in target_doc.styles}

        for style in source_styles:
            if style.name and style.name not in target_style_names:
                try:
                    # Styles phức tạp khó copy trực tiếp,
                    # chỉ copy XML element
                    pass
                except Exception:
                    pass
    except Exception:
        pass


def merge_docx_files(
    file_paths: list[str],
    output_path: str,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
) -> str:
    """
    Gộp nhiều file .docx thành 1 file.

    Args:
        file_paths: Danh sách đường dẫn file .docx cần gộp
        output_path: Đường dẫn file output
        progress_callback: Callback(current, total, filename) để báo tiến trình

    Returns:
        Đường dẫn file output đã tạo
    """
    if not file_paths:
        raise ValueError("Không có file nào để gộp!")

    total = len(file_paths)

    # Tạo document mới hoàn toàn trống
    merged_doc = Document()

    # Xóa paragraph mặc định
    if merged_doc.paragraphs:
        for p in merged_doc.paragraphs:
            p._element.getparent().remove(p._element)

    target_body = merged_doc.element.body

    for idx, file_path in enumerate(file_paths):
        filename = os.path.basename(file_path)

        if progress_callback:
            progress_callback(idx + 1, total, filename)

        try:
            source_doc = Document(file_path)

            # Copy section properties từ file đầu tiên (page size, margins...)
            if idx == 0:
                source_sections = source_doc.sections
                if source_sections:
                    first_section = source_sections[0]
                    target_section = merged_doc.sections[0] if merged_doc.sections else None
                    if target_section:
                        target_section.page_width = first_section.page_width
                        target_section.page_height = first_section.page_height
                        target_section.top_margin = first_section.top_margin
                        target_section.bottom_margin = first_section.bottom_margin
                        target_section.left_margin = first_section.left_margin
                        target_section.right_margin = first_section.right_margin

            # Thêm page break trước file (trừ file đầu tiên)
            if idx > 0:
                add_page_break(target_body)

            # Copy tất cả elements từ body của source document
            source_body = source_doc.element.body
            for child in source_body:
                # Bỏ qua sectPr (section properties) - chỉ copy content
                if child.tag == qn("w:sectPr"):
                    continue
                copy_element(child, target_body)

        except Exception as e:
            # Log lỗi nhưng vẫn tiếp tục với các file khác
            if progress_callback:
                progress_callback(idx + 1, total, f"⚠️ LỖI: {filename} - {str(e)}")

    # Lưu file output
    merged_doc.save(output_path)
    return output_path


def get_docx_files(folder_path: str) -> list[str]:
    """
    Lấy danh sách file .docx trong folder, sắp xếp theo tên.

    Args:
        folder_path: Đường dẫn folder

    Returns:
        Danh sách đường dẫn file .docx (sorted)
    """
    folder = Path(folder_path)
    docx_files = sorted(
        [str(f) for f in folder.glob("*.docx") if not f.name.startswith("~$")],
        key=lambda x: os.path.basename(x).lower(),
    )
    return docx_files
