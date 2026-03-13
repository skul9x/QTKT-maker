"""
DOCX Generator for QTKT documents
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import re

from config import (
    FONT_NAME, FONT_SIZE, LINE_SPACING,
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
    PARAGRAPH_BEFORE, PARAGRAPH_AFTER
)


class DocxGenerator:
    """Generator để tạo file DOCX theo format chuẩn BYT"""
    
    def __init__(self):
        self.document = None
    
    def create_document(self, content: dict, output_path: str):
        """
        Tạo file DOCX từ nội dung QTKT
        
        Args:
            content: dict chứa ten_quy_trinh và noi_dung
            output_path: đường dẫn file output
        """
        self.document = Document()
        
        # Setup page margins
        self._setup_page_format()
        
        # Add title
        self._add_title(content.get("ten_quy_trinh", "QUY TRÌNH KỸ THUẬT"))
        
        # Add content
        self._add_content(content.get("noi_dung", ""))
        
        # Save
        self.document.save(output_path)
        
        return output_path
    
    def _setup_page_format(self):
        """Thiết lập format trang theo QĐ 3023/QĐ-BYT"""
        section = self.document.sections[0]
        
        # Margins
        section.top_margin = Cm(MARGIN_TOP)
        section.bottom_margin = Cm(MARGIN_BOTTOM)
        section.left_margin = Cm(MARGIN_LEFT)
        section.right_margin = Cm(MARGIN_RIGHT)
        
        # Header for page numbers
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_title(self, title: str):
        """Thêm tiêu đề quy trình"""
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = para.add_run(title.upper())
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE + 1)
        
        # Add spacing after title
        para.paragraph_format.space_after = Pt(12)
    
    # Danh sách các pattern cần in đậm (theo yêu cầu user)
    BOLD_PATTERNS = [
        r"^TÊN KỸ THUẬT[:;]?",
        r"^\d+\.\s+ĐẠI CƯƠNG",
        r"^\d+\.\s+CHỈ ĐỊNH",
        r"^\d+\.\s+CHỐNG CHỈ ĐỊNH",
        r"^\d+\.\s+THẬN TRỌNG",
        r"^\d+\.\s+CHUẨN BỊ",
        r"^\d+\.\d+\.\s+Người thực hiện",
        r"^\d+\.\d+\.\s+Thuốc",
        r"^\d+\.\d+\.\s+Vật tư",
        r"^\d+\.\d+\.\s+Trang thiết bị",
        r"^\d+\.\d+\.\s+Người bệnh",
        r"^\d+\.\d+\.\s+Hồ sơ bệnh án",
        r"^\d+\.\d+\.\s+Thời gian thực hiện kỹ thuật",
        r"^\d+\.\d+\.\s+Địa điểm thực hiện kỹ thuật",
        r"^\d+\.\d+\.\s+Kiểm tra hồ sơ",
        r"^\d+\.\s+TIẾN HÀNH QUY TRÌNH KỸ THUẬT", # Added section 6 just in case
        r"^\d+\.\d+\.\s+Bước\s+\d+", # Bold headers like "6.1. Bước 1:", "6.2. Bước 2:"
        r"^\d+\.\d+\.\s+Kết thúc quy trình", # Bold header "6.x. Kết thúc quy trình"
        r"^\d+\.\s+THEO DÕI VÀ XỬ TRÍ TAI BIẾN",
        r"^TÀI LIỆU THAM KHẢO"
    ]

    def _add_content(self, content: str):
        """Thêm nội dung QTKT"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            para = self.document.add_paragraph()
            
            # Check if line matches any bold pattern
            is_bold_header = any(re.match(pattern, line, re.IGNORECASE) for pattern in self.BOLD_PATTERNS)
            
            # Check if it's a generic heading (starts with number or ##) if not specialized bold
            is_generic_heading = bool(re.match(r'^(\d+\.|\#{1,3})\s', line))
            
            # Clean markdown symbols
            line = re.sub(r'^\#{1,3}\s*', '', line)
            line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # Bold markdown
            line = re.sub(r'\*([^*]+)\*', r'\1', line)  # Italic markdown
            
            run = para.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            
            if is_bold_header or is_generic_heading:
                run.bold = True
            
            # Paragraph formatting
            para.paragraph_format.space_before = Pt(PARAGRAPH_BEFORE)
            para.paragraph_format.space_after = Pt(PARAGRAPH_AFTER)
            para.paragraph_format.line_spacing = LINE_SPACING
    
    def preview_content(self, content: dict) -> str:
        """
        Tạo preview text từ content
        
        Args:
            content: dict chứa noi_dung
            
        Returns:
            str: preview text
        """
        text = content.get("noi_dung", "")
        # Clean markdown for preview
        text = re.sub(r'\#{1,3}\s*', '', text)
        return text
